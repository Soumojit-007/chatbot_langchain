
# import streamlit as st
# from dotenv import load_dotenv
# from PyPDF2 import PdfReader
# import docx
# import io
# import os
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.text_splitter import CharacterTextSplitter
# from langchain.chains.question_answering import load_qa_chain
# from langchain.docstore.document import Document

# # Correct import for summarization
# from langchain.chains.summarize import load_summarize_chain
# from langchain.prompts import PromptTemplate

# # Load environment variables
# load_dotenv()
# api_key = os.getenv("GOOGLE_API_KEY")

# # Initialize the LLM
# llm = ChatGoogleGenerativeAI(
#     model="gemini-1.5-flash",
#     google_api_key=api_key
# )

# # Streamlit UI
# st.title("📄 Google Gemini Chatbot - LangChain")

# uploaded_file = st.file_uploader("Upload a PDF or Word Document", type=["pdf", "docx"])

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# if uploaded_file is not None:
#     try:
#         # Extract text from file
#         text = ""
#         if uploaded_file.name.endswith(".pdf"):
#             pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
#             for page in pdf_reader.pages:
#                 text += page.extract_text() or ""
#         elif uploaded_file.name.endswith(".docx"):
#             text = extract_text_from_docx(uploaded_file)

#         # Split into chunks
#         text_splitter = CharacterTextSplitter(separator="\n", chunk_size=1000, chunk_overlap=100)
#         texts = text_splitter.split_text(text)
#         docs = [Document(page_content=chunk) for chunk in texts]

#         # QA Chain (unchanged)
#         qa_chain = load_qa_chain(llm, chain_type="stuff")

#         # Custom prompt for summarization
#         custom_prompt = PromptTemplate(
#             input_variables=["text"],
#             template="Summarize the following document focusing on the main points:\n\n{text}\n\nSummary:"
#         )

#         # Summarization chain with custom prompt
#         summarization_chain = load_summarize_chain(
#             llm, chain_type="stuff",
#             prompt=custom_prompt
#         )

#         # Question input and response
#         query = st.text_input("Ask a question about the document:")
#         if query:
#             response = qa_chain.invoke({
#                 "input_documents": docs,
#                 "question": query
#             })
#             st.subheader("Response:")
#             st.write(response["output_text"])

#         # Summarize button
#         if st.button("Summarize Document"):
#             summary = summarization_chain.run(docs)
#             st.subheader("📌 Summary:")
#             st.write(summary)

#     except Exception as e:
#         st.error(f"❌ Failed to process the file: {e}")

import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
import io
import os
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.question_answering import load_qa_chain
from langchain.docstore.document import Document
from langchain.chains.summarize import load_summarize_chain
from langchain.prompts import PromptTemplate
import time
from datetime import datetime
import traceback

# Load environment variables
load_dotenv()

# Get API key from environment or Streamlit secrets
try:
    api_key = st.secrets["GOOGLE_API_KEY"] if "GOOGLE_API_KEY" in st.secrets else os.getenv("GOOGLE_API_KEY")
except:
    api_key = os.getenv("GOOGLE_API_KEY")

# Page configuration
st.set_page_config(
    page_title="📄 AI Document Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for advanced styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    
    .feature-card {
        background: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    
    .stats-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        border-left: 4px solid #667eea;
        background: #f1f3f4;
    }
    
    .success-message {
        background: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 10px;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    
    .sidebar-content {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'document_processed' not in st.session_state:
    st.session_state.document_processed = False
if 'document_stats' not in st.session_state:
    st.session_state.document_stats = {}
if 'docs' not in st.session_state:
    st.session_state.docs = []
if 'qa_chain' not in st.session_state:
    st.session_state.qa_chain = None
if 'summarization_chain' not in st.session_state:
    st.session_state.summarization_chain = None
if 'llm_initialized' not in st.session_state:
    st.session_state.llm_initialized = False

# Initialize the LLM with error handling
def initialize_llm():
    try:
        if not api_key:
            st.error("❌ Google API Key not found. Please check your environment variables or Streamlit secrets.")
            return None
        
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=api_key,
            temperature=0.3,
            max_tokens=2048,
            timeout=60  # Add timeout for deployment
        )
        
        # Test the LLM with a simple query
        test_response = llm.invoke("Hello")
        st.session_state.llm_initialized = True
        return llm
    except Exception as e:
        st.error(f"❌ Failed to initialize LLM: {str(e)}")
        st.session_state.llm_initialized = False
        return None

# Initialize LLM only once
if not st.session_state.llm_initialized:
    with st.spinner("🔄 Initializing AI model..."):
        llm = initialize_llm()
else:
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        google_api_key=api_key,
        temperature=0.3,
        max_tokens=2048,
        timeout=60
    )

# Sidebar
with st.sidebar:
    st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    st.title("🔧 Control Panel")
    
    # System Status
    st.subheader("🔍 System Status")
    llm_status = "✅ Connected" if st.session_state.llm_initialized else "❌ Connection Failed"
    st.markdown(f"**LLM Status:** {llm_status}")
    st.markdown(f"**Session Time:** {datetime.now().strftime('%H:%M:%S')}")
    
    # Document Statistics
    if st.session_state.document_processed:
        st.subheader("📊 Document Statistics")
        stats = st.session_state.document_stats
        st.metric("📄 Total Pages", stats.get('pages', 'N/A'))
        st.metric("📝 Word Count", stats.get('words', 'N/A'))
        st.metric("🔤 Character Count", stats.get('characters', 'N/A'))
        st.metric("📦 Text Chunks", stats.get('chunks', 'N/A'))
    
    # Settings
    st.subheader("⚙️ Settings")
    chunk_size = st.slider("Chunk Size", 500, 2000, 1000, 100)
    chunk_overlap = st.slider("Chunk Overlap", 50, 200, 100, 25)
    
    # Clear chat history
    if st.button("🗑️ Clear Chat History", use_container_width=True):
        st.session_state.chat_history = []
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)

# Main header
st.markdown("""
<div class="main-header">
    <h1>🤖 AI Document Assistant</h1>
    <p>Upload, Analyze, and Chat with your Documents using Google Gemini</p>
</div>
""", unsafe_allow_html=True)

# Check if LLM is properly initialized
if not st.session_state.llm_initialized:
    st.error("⚠️ AI model not initialized. Please check your API key configuration.")
    st.info("💡 Make sure to add your GOOGLE_API_KEY to Streamlit secrets or environment variables.")
    st.stop()

# Feature overview
col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div class="feature-card">
        <h4>📄 Document Upload</h4>
        <p>Support for PDF and Word documents with intelligent text extraction</p>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div class="feature-card">
        <h4>💬 Smart Chat</h4>
        <p>Ask questions and get contextual answers from your documents</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div class="feature-card">
        <h4>📝 Auto Summary</h4>
        <p>Generate comprehensive summaries with key insights</p>
    </div>
    """, unsafe_allow_html=True)

# File upload section
st.subheader("📁 Document Upload")
uploaded_file = st.file_uploader(
    "Choose a PDF or Word document",
    type=["pdf", "docx"],
    help="Upload your document to start chatting with it!"
)

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""

def calculate_stats(text, chunks):
    words = len(text.split())
    characters = len(text)
    return {
        'words': words,
        'characters': characters,
        'chunks': len(chunks)
    }

if uploaded_file is not None:
    try:
        with st.spinner("🔄 Processing document... Please wait"):
            # Extract text from file
            text = ""
            if uploaded_file.name.endswith(".pdf"):
                try:
                    pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
                    pages = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                except Exception as e:
                    st.error(f"Error reading PDF: {str(e)}")
                    st.stop()
            elif uploaded_file.name.endswith(".docx"):
                text = extract_text_from_docx(uploaded_file)
                pages = "N/A"
            
            if not text.strip():
                st.error("❌ No text could be extracted from the document. Please check the file.")
                st.stop()

            # Split into chunks with error handling
            try:
                text_splitter = CharacterTextSplitter(
                    separator="\n", 
                    chunk_size=chunk_size, 
                    chunk_overlap=chunk_overlap
                )
                texts = text_splitter.split_text(text)
                
                if not texts:
                    st.error("❌ Could not split document into chunks. Please try a different document.")
                    st.stop()
                
                docs = [Document(page_content=chunk) for chunk in texts]
                st.session_state.docs = docs
                
            except Exception as e:
                st.error(f"Error splitting text: {str(e)}")
                st.stop()

            # Calculate statistics
            stats = calculate_stats(text, texts)
            if pages != "N/A":
                stats['pages'] = pages
            st.session_state.document_stats = stats
            st.session_state.document_processed = True

            # Initialize chains with error handling
            try:
                # QA Chain with timeout handling
                st.session_state.qa_chain = load_qa_chain(llm, chain_type="stuff")

                # Custom prompt for summarization
                custom_prompt = PromptTemplate(
                    input_variables=["text"],
                    template="Summarize the following document focusing on the main points:\n\n{text}\n\nSummary:"
                )

                # Summarization chain with custom prompt
                st.session_state.summarization_chain = load_summarize_chain(
                    llm, chain_type="stuff",
                    prompt=custom_prompt
                )
            except Exception as e:
                st.error(f"Error initializing chains: {str(e)}")
                st.stop()

        # Success message
        st.markdown(f"""
        <div class="success-message">
            <h4>✅ Document Successfully Processed!</h4>
            <p><strong>File:</strong> {uploaded_file.name}</p>
            <p><strong>Size:</strong> {len(text)} characters | <strong>Chunks:</strong> {len(texts)}</p>
        </div>
        """, unsafe_allow_html=True)

        # Main interface tabs
        tab1, tab2, tab3 = st.tabs(["💬 Chat with Document", "📝 Generate Summary", "📊 Document Analysis"])

        with tab1:
            st.subheader("💬 Ask Questions About Your Document")
            
            # Chat interface
            query = st.text_input(
                "Type your question here...",
                placeholder="e.g., What are the main topics discussed in this document?",
                key="chat_input"
            )
            
            col1, col2 = st.columns([1, 4])
            with col1:
                ask_button = st.button("🚀 Ask Question", use_container_width=True)
            
            if ask_button and query and st.session_state.qa_chain:
                try:
                    with st.spinner("🤔 Thinking... Please wait"):
                        start_time = time.time()
                        
                        # Add timeout handling for the query
                        response = st.session_state.qa_chain.invoke({
                            "input_documents": st.session_state.docs,
                            "question": query
                        })
                        
                        end_time = time.time()
                        
                        # Add to chat history
                        st.session_state.chat_history.append({
                            "question": query,
                            "answer": response["output_text"],
                            "timestamp": datetime.now().strftime("%H:%M:%S"),
                            "processing_time": f"{end_time - start_time:.2f}s"
                        })
                    
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"❌ Error processing question: {str(e)}")
                    st.info("💡 Try rephrasing your question or check your internet connection.")
            
            # Display chat history
            if st.session_state.chat_history:
                st.subheader("💭 Conversation History")
                for i, chat in enumerate(reversed(st.session_state.chat_history)):
                    with st.expander(f"Q{len(st.session_state.chat_history)-i}: {chat['question'][:50]}...", expanded=(i==0)):
                        st.markdown(f"**⏰ Time:** {chat['timestamp']} | **⚡ Processing:** {chat['processing_time']}")
                        st.markdown(f"**❓ Question:** {chat['question']}")
                        st.markdown(f"**✨ Answer:** {chat['answer']}")

        with tab2:
            st.subheader("📝 Document Summary")
            
            # Summary options
            col1, col2 = st.columns(2)
            with col1:
                summary_type = st.selectbox(
                    "Summary Type",
                    ["Comprehensive", "Key Points", "Executive Summary", "Technical Overview"]
                )
            
            if st.button("📋 Generate Summary", use_container_width=True):
                if st.session_state.summarization_chain:
                    try:
                        with st.spinner("📝 Generating summary... This may take a moment"):
                            start_time = time.time()
                            
                            # Add timeout and error handling for summarization
                            summary = st.session_state.summarization_chain.run(st.session_state.docs)
                            
                            end_time = time.time()
                            
                            st.markdown("### 📌 Document Summary")
                            st.markdown(f"**Summary Type:** {summary_type}")
                            st.markdown(f"**Generated in:** {end_time - start_time:.2f} seconds")
                            st.markdown("---")
                            st.write(summary)
                            
                            # Download summary option
                            st.download_button(
                                label="📥 Download Summary",
                                data=summary,
                                file_name=f"summary_{uploaded_file.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                                mime="text/plain"
                            )
                    except Exception as e:
                        st.error(f"❌ Error generating summary: {str(e)}")
                        st.info("💡 Try again or check your internet connection.")

        with tab3:
            st.subheader("📊 Document Analysis")
            
            # Display detailed statistics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.markdown(f"""
                <div class="stats-card">
                    <h3>{stats.get('pages', 'N/A')}</h3>
                    <p>Pages</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                <div class="stats-card">
                    <h3>{stats['words']:,}</h3>
                    <p>Words</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col3:
                st.markdown(f"""
                <div class="stats-card">
                    <h3>{stats['characters']:,}</h3>
                    <p>Characters</p>
                </div>
                """, unsafe_allow_html=True)
            
            with col4:
                st.markdown(f"""
                <div class="stats-card">
                    <h3>{stats['chunks']}</h3>
                    <p>Text Chunks</p>
                </div>
                """, unsafe_allow_html=True)
            
            # Additional analysis
            st.subheader("📈 Text Analysis")
            
            # Sample text preview
            with st.expander("📖 Document Preview (First 500 characters)"):
                st.text(text[:500] + "..." if len(text) > 500 else text)
            
            # Chunk information
            with st.expander("🔍 Chunk Analysis"):
                chunk_sizes = [len(chunk) for chunk in texts]
                avg_chunk_size = sum(chunk_sizes) / len(chunk_sizes)
                st.metric("Average Chunk Size", f"{avg_chunk_size:.0f} characters")
                st.metric("Largest Chunk", f"{max(chunk_sizes)} characters")
                st.metric("Smallest Chunk", f"{min(chunk_sizes)} characters")

    except Exception as e:
        st.error(f"❌ Failed to process the file: {str(e)}")
        st.error(f"Debug info: {traceback.format_exc()}")
        st.info("💡 Please make sure your file is not corrupted and try again.")

else:
    # Welcome message
    st.markdown("""
    ### 👋 Welcome to AI Document Assistant!
    
    **Getting Started:**
    1. 📤 Upload a PDF or Word document using the file uploader above
    2. 💬 Ask questions about your document content
    3. 📝 Generate automatic summaries
    4. 📊 View detailed document analysis
    
    **Supported Formats:** PDF (.pdf), Word Documents (.docx)
    """)
    
    # Example questions
    with st.expander("💡 Example Questions You Can Ask"):
        st.markdown("""
        - What are the main topics discussed in this document?
        - Can you summarize the key findings?
        - What are the conclusions or recommendations?
        - Who are the main people or organizations mentioned?
        - What dates or events are referenced?
        - What are the technical specifications mentioned?
        """)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 1rem;">
    🤖 Powered by Google Gemini & LangChain | Built with Streamlit<br>
    <small>For best results, use clear and specific questions about your document content.</small>
</div>
""", unsafe_allow_html=True)